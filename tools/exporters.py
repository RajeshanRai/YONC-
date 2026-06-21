from io import BytesIO
import json
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


def sanitize_filename(value: str) -> str:
    cleaned = re.sub(r'[^A-Za-z0-9_-]+', '_', (value or '').strip())
    return cleaned[:80] or 'document'


def _load_payload(content):
    if isinstance(content, dict):
        return content
    if not content:
        return {}
    if isinstance(content, str):
        text = content.strip()
        if text.startswith('{'):
            try:
                loaded = json.loads(text)
                if isinstance(loaded, dict):
                    return loaded
            except Exception:
                pass
        return {'text': content}
    return {'text': str(content)}


def _docx_styles(document: Document):
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)

    normal = document.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10.5)


def _docx_add_section_rule(paragraph):
    """Add a bottom border line under a heading paragraph for a clean section divider."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0F172A')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _docx_add_heading(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    _docx_add_section_rule(paragraph)


def _docx_add_body(document: Document, text: str, *, italic: bool = False, bold: bool = False, bullet: bool = False):
    paragraph = document.add_paragraph(style='List Bullet' if bullet else None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)


def _docx_render_from_payload(document: Document, payload: dict):
    full_name = (payload.get('full_name') or payload.get('name') or '').strip()
    contact_line = (payload.get('contact_line') or '').strip()
    summary = (payload.get('summary') or '').strip()
    skills = [str(item).strip() for item in payload.get('skills', []) if str(item).strip()]
    experience_bullets = [str(item).strip() for item in payload.get('experience_bullets', []) if str(item).strip()]
    education = (payload.get('education') or '').strip()
    achievements = (payload.get('achievements') or '').strip()
    certifications = (payload.get('certifications') or '').strip()
    projects = (payload.get('projects') or '').strip()
    languages = (payload.get('languages') or '').strip()
    references = (payload.get('references') or '').strip()
    additional_information = (payload.get('additional_information') or '').strip()
    paragraphs = [str(item).strip() for item in payload.get('paragraphs', []) if str(item).strip()]

    if full_name:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(2)
        run = title.add_run(full_name)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    job_title_display = (payload.get('job_title') or '').strip()
    if job_title_display:
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(2)
        run = subtitle.add_run(job_title_display)
        run.italic = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    if contact_line:
        contact = document.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.paragraph_format.space_after = Pt(6)
        run = contact.add_run(contact_line)
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    if paragraphs:
        for paragraph_text in paragraphs:
            _docx_add_body(document, paragraph_text)
        return

    if summary:
        _docx_add_heading(document, 'Professional Summary')
        _docx_add_body(document, summary)

    if skills:
        _docx_add_heading(document, 'Core Competencies')
        _docx_add_body(document, ', '.join(skills))

    if experience_bullets:
        _docx_add_heading(document, 'Professional Experience')
        for bullet in experience_bullets:
            _docx_add_body(document, bullet, bullet=True)

    if education:
        _docx_add_heading(document, 'Education')
        _docx_add_body(document, education)

    if achievements:
        _docx_add_heading(document, 'Achievements')
        for line in (achievements.splitlines() if '\n' in achievements else [achievements]):
            line = line.strip()
            if line:
                _docx_add_body(document, line, bullet=True)

    extra_sections = [
        ('Certifications', certifications),
        ('Projects', projects),
        ('Languages', languages),
        ('References', references),
        ('Additional Information', additional_information),
    ]

    for heading, content in extra_sections:
        if content:
            _docx_add_heading(document, heading)
            for line in (content.splitlines() if '\n' in content else [content]):
                line = line.strip()
                if line:
                    _docx_add_body(document, line, bullet=True)


def build_docx_bytes(content) -> bytes:
    payload = _load_payload(content)
    document = Document()
    _docx_styles(document)

    if 'text' in payload and len(payload) == 1:
        for line in str(payload['text']).splitlines():
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(line)
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
    else:
        _docx_render_from_payload(document, payload)

    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream.read()


def _pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='ResumeTitle',
        parent=styles['Title'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#0f172a'),
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name='ResumeSubtitle',
        parent=styles['BodyText'],
        fontName='Helvetica-Oblique',
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#334155'),
        spaceAfter=3,
    ))
    styles.add(ParagraphStyle(
        name='ResumeContact',
        parent=styles['BodyText'],
        fontName='Helvetica-Oblique',
        fontSize=9.5,
        leading=12,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#334155'),
        spaceAfter=10,
    ))
    styles.add(ParagraphStyle(
        name='ResumeHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=13,
        alignment=TA_LEFT,
        textColor=colors.HexColor('#0f172a'),
        spaceBefore=12,
        spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        name='ResumeBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10.2,
        leading=13,
        alignment=TA_JUSTIFY,
        textColor=colors.HexColor('#1f2937'),
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name='ResumeBullet',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10.2,
        leading=13,
        leftIndent=12,
        firstLineIndent=-10,
        bulletIndent=0,
        alignment=TA_JUSTIFY,
        textColor=colors.HexColor('#1f2937'),
        spaceAfter=4,
    ))
    return styles


def _pdf_render_story(payload: dict):
    styles = _pdf_styles()
    story = []

    full_name = (payload.get('full_name') or payload.get('name') or '').strip()
    job_title = (payload.get('job_title') or '').strip()
    contact_line = (payload.get('contact_line') or '').strip()
    summary = (payload.get('summary') or '').strip()
    skills = [str(item).strip() for item in payload.get('skills', []) if str(item).strip()]
    experience_bullets = [str(item).strip() for item in payload.get('experience_bullets', []) if str(item).strip()]
    education = (payload.get('education') or '').strip()
    achievements = (payload.get('achievements') or '').strip()
    certifications = (payload.get('certifications') or '').strip()
    projects = (payload.get('projects') or '').strip()
    languages = (payload.get('languages') or '').strip()
    references = (payload.get('references') or '').strip()
    additional_information = (payload.get('additional_information') or '').strip()
    paragraphs = [str(item).strip() for item in payload.get('paragraphs', []) if str(item).strip()]

    def _section_heading(label):
        story.append(Paragraph(label.upper(), styles['ResumeHeading']))
        story.append(HRFlowable(
            width='100%', thickness=1,
            color=colors.HexColor('#0f172a'),
            spaceBefore=2, spaceAfter=6,
        ))

    if full_name:
        story.append(Paragraph(full_name, styles['ResumeTitle']))
    if job_title:
        story.append(Paragraph(job_title, styles['ResumeSubtitle']))
    if contact_line:
        story.append(Paragraph(contact_line, styles['ResumeContact']))

    if paragraphs:
        for paragraph_text in paragraphs:
            story.append(Paragraph(paragraph_text, styles['ResumeBody']))
            story.append(Spacer(1, 4))
        return story

    if summary:
        _section_heading('Professional Summary')
        story.append(Paragraph(summary, styles['ResumeBody']))

    if skills:
        _section_heading('Core Competencies')
        story.append(Paragraph(' \u2022 '.join(skills), styles['ResumeBody']))

    if experience_bullets:
        _section_heading('Professional Experience')
        bullet_flow = []
        for bullet in experience_bullets:
            bullet_flow.append(ListItem(Paragraph(bullet, styles['ResumeBullet'])))
        story.append(ListFlowable(bullet_flow, bulletType='bullet', leftIndent=12))

    if education:
        _section_heading('Education')
        story.append(Paragraph(education, styles['ResumeBody']))

    if achievements:
        _section_heading('Achievements')
        achievement_items = [line.strip() for line in achievements.splitlines() if line.strip()] or [achievements]
        if len(achievement_items) > 1:
            bullet_flow = [ListItem(Paragraph(item, styles['ResumeBullet'])) for item in achievement_items]
            story.append(ListFlowable(bullet_flow, bulletType='bullet', leftIndent=12))
        else:
            story.append(Paragraph(achievements, styles['ResumeBody']))

    extra_sections = [
        ('Certifications', certifications),
        ('Projects', projects),
        ('Languages', languages),
        ('References', references),
        ('Additional Information', additional_information),
    ]

    for heading, content in extra_sections:
        if content:
            _section_heading(heading)
            items = [line.strip() for line in content.splitlines() if line.strip()] or [content]
            if len(items) > 1:
                bullet_flow = [ListItem(Paragraph(item, styles['ResumeBullet'])) for item in items]
                story.append(ListFlowable(bullet_flow, bulletType='bullet', leftIndent=12))
            else:
                story.append(Paragraph(content, styles['ResumeBody']))

    return story


def build_pdf_bytes(content) -> bytes:
    payload = _load_payload(content)
    stream = BytesIO()
    document = SimpleDocTemplate(
        stream,
        pagesize=A4,
        rightMargin=48,
        leftMargin=48,
        topMargin=48,
        bottomMargin=48,
    )
    story = _pdf_render_story(payload)
    if not story:
        story = [Paragraph(' ', _pdf_styles()['ResumeBody'])]
    document.build(story)
    stream.seek(0)
    return stream.read()